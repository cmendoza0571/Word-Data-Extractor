import docx

#Notes
#no page numbers
#no line numbers for images
#no way to access charts(?)
#Need to find way to determine where an actual paragraph begins and ends

doc = docx.Document("C:/Users/Luis/Desktop/Projects/TestWordFile.docx") 
#change filepath to where you have docx saved

all_paras = doc.paragraphs #all_paras = all paragraphs/lines of the docx file

#print(len(all_paras)) #print out how many paragraphs docx file has

boldTextStrings = [] #testing out printing from array of bold text from docx file, ignore this
italicText = "" #String of all italic text from docx file, ignore this
fontsUsed = [] #Array of fonts used in docx file

#For loop for going through all the text in a docx file
for para in all_paras: #goes through each paragraph/line in the docx
    for run in para.runs: #goes through each run in a paragraph/line
        if run.font.name not in fontsUsed:
                fontsUsed.append(run.font.name)  #add font to fontsUsed array if not already in array
        # elif run.bold:
        #     boldTextStrings.append(run.text) 
        #     if(run.text == '.' or run.text == ',' or run.text == '!')
        # elif run.italic: 
        #     italicText += run.text
        elif para.style.name == 'List Paragraph':
            pass #Stops lists from printing because the for loop prints out each item in a list multiple times
        elif para.text == "": 
            pass #Ignores blank lines
        else:
            print(para.text) #Print out each line of text
            print("---------------") #Just a way of separating lines in the output


#Print out array of fonts used
print("Fonts used in this document:",fontsUsed)


#for loop for printing out items in a list
#Need to find out a way to get the list number
for para in all_paras:
    if para.style.name == 'List Paragraph':
            print(para.text)
            

#For loop for printing out the text in a table
for table in doc.tables: 
    rowIndex = 0
    colIndex = 0
    for row in table.rows:
        rowIndex = rowIndex + 1
        for cell in row.cells:
            if colIndex >= len(table.columns):
                colIndex = 0
            colIndex = colIndex + 1
            for para in cell.paragraphs:
                print("Row ", rowIndex, "Column ", colIndex, ": ", para.text)

for image in doc.inline_shapes:
    print ("There was a picture here, figure out a way to find out what line it was on")
    #print(image.image_path) image.image_path is not a real thing, need to find something like it

#print(boldTextStrings) #testing out printing text from run.text as an array


#Printing out bold, italic, and underlined text
print("Bold Text: ")
for para in all_paras:
    for run in para.runs:
        if run.bold:
            print(run.text)
print("Italicized Text: ")
for para in all_paras:
    for run in para.runs:
        if run.italic:
            print(run.text)
print("Underlined Text: ")
for para in all_paras:
    for run in para.runs:
        if run.underline:
            print(run.text)

#Printing out different types of headers and footers
firstPageHeader = doc.sections[0].first_page_header 
for para in firstPageHeader.paragraphs:
    print(para.text)
    
firstPageFooter = doc.sections[0].first_page_footer
for para in firstPageFooter.paragraphs:
    print(para.text)
            
header = doc.sections[0].header
for para in header.paragraphs:
    print(para.text)

footer = doc.sections[0].footer
for para in footer.paragraphs:
    print(para.text)
    
#other types of headers and footers    
#even_page_footer
#even_page_header    